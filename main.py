import tkinter as tk
import openpyxl
import random
import os
import datetime
import docx
import sys
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
def calculate_bmi():
    try:
        # Kullanıcının girdiği değerleri al
        ad_soyad = entry_ad_soyad.get()
        meal_number = int(entry_ogun_sayisi.get())
        kilo = float(entry_kilo.get())
        boy = float(entry_boy.get())
        yas = int(entry_yas.get())

        # klasorü ve diyeti bulmak için yol
        klasor_bul = "DIETS"

        # Adınızı ve soyadınızı giriniz

        # Kaç öğün yiyorsunuz
        if meal_number <= 2:
            klasor_bul += "\\TWO_MEAL_DIETS\\"
        elif meal_number >= 3:
            klasor_bul += "\\THREE_MEAL_DIETS\\"
        else:
            print("Geçersiz bir öğün sayısı girdiniz. Lütfen botu kapatıp tekrar açın, 2 veya 3 girin.")

            # Klasörleri listeleyin ve Rastgele bir klasör seçin
        aylar = [f for f in os.listdir(klasor_bul) if os.path.isdir(os.path.join(klasor_bul, f))]
        klasor_bul += random.choice(aylar)


        # BMİ HESAPLA:
        bmi = kilo / (boy ** 2)
        if bmi < 21:
            klasor_bul += "\\21_25bmi"
        elif bmi >= 21 and bmi <= 25:
            klasor_bul += "\\21_25bmi"
        elif bmi >= 26 and bmi <= 29:
            klasor_bul += "\\26_29bmi"
        elif bmi >= 30 and bmi <= 33:
            klasor_bul += "\\30_33bmi"
        elif bmi >= 34 and bmi <= 37:
            klasor_bul += "\\34_37bmi"
        elif bmi > 37:
            klasor_bul += "\\34_37bmi"

        # Geçmemeniz Gereken Kilo:
        gecmemesi_gereken_kilo = (25 * boy * boy) // 1

        # İdael Kilo:
        ideal_kilo = (21 * boy ** 2) // 1

        # Kontrol Tarihiniz:
        gecen_zaman = datetime.timedelta(days=7)
        simdiki_zaman = datetime.datetime.now()
        tarih = simdiki_zaman.strftime("%m/%d/%Y")
        sonraki_zaman = simdiki_zaman + gecen_zaman
        ayın_kacı = datetime.datetime.now()

        # raskele hafta
        diyet_dosyası = [f for f in os.listdir(klasor_bul) if f.endswith(".docx")]
        hafta = random.choice(diyet_dosyası)
        klasor_bul += ("\\" + str(hafta))

        # Diyeti bul ve ilk sayfadaki bilgileri yerleştir
        diyet = docx.Document(klasor_bul)

        # ADI SOYADI
        ad_bilgisi = diyet.paragraphs[0]
        ad_bilgisi.text = "ADI SOYADI: " + str(ad_soyad)
        ad_bilgisi.runs[0].bold = True
        run_ad_bilgisi = ad_bilgisi.runs[0]
        run_ad_bilgisi.font.name = 'Comic Sans MS'
        run_ad_bilgisi.font.size = Pt(14)

        # Ağırlık
        kilo_boy_yas_bilgisi = diyet.paragraphs[3]
        kilo_boy_yas_bilgisi.text = "Ağırlık: " + str(kilo) + "             Boy:" + str(
            boy) + "                Yaş: " + str(yas)
        kilo_boy_yas_bilgisi.runs[0].bold = True
        kilo_boy_yas_bilgisi.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
        run_kilo_boy_yas_bilgisi = kilo_boy_yas_bilgisi.runs[0]
        run_kilo_boy_yas_bilgisi.font.name = 'Comic Sans MS'
        run_kilo_boy_yas_bilgisi.font.size = Pt(14)

        # Hedeflenen kilo kaybı
        h_k_bilgisi = diyet.paragraphs[5]
        h_k_bilgisi.text = "Hedeflenen kilo kaybı: 1 kg (7 Günde)   " + "Kontrol Tarihi: " + str(
            sonraki_zaman.strftime("%m/%d/%Y"))
        h_k_bilgisi.runs[0].bold = True
        run_h_k_bilgisi = h_k_bilgisi.runs[0]
        run_h_k_bilgisi.font.name = 'Comic Sans MS'
        run_h_k_bilgisi.font.size = Pt(13)

        # BKI
        bmi = bmi // 1
        bki_bilgisi = diyet.paragraphs[9]
        bki_bilgisi.text = "BKI= " + str(bmi)
        bki_bilgisi.runs[0].bold = False
        bki_bilgisi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bki_bilgisi = bki_bilgisi.runs[0]
        run_bki_bilgisi.font.name = 'Comic Sans MS'
        run_bki_bilgisi.font.size = Pt(14)

        # İdeal kilonuz
        i_k_bilgisi = diyet.paragraphs[10]
        i_k_bilgisi.text = "İdeal kilonuz= " + str(ideal_kilo) + " kg"
        i_k_bilgisi.runs[0].bold = False
        i_k_bilgisi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_i_k_bilgisi = i_k_bilgisi.runs[0]
        run_i_k_bilgisi.font.name = 'Comic Sans MS'
        run_i_k_bilgisi.font.size = Pt(14)

        # Geçmemeniz gereken kilo
        g_c_bilgisi = diyet.paragraphs[11]
        g_c_bilgisi.text = "Geçmemeniz gereken kilo= " + str(gecmemesi_gereken_kilo) + " kg"
        g_c_bilgisi.runs[0].bold = False
        g_c_bilgisi.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_g_c_bilgisi = g_c_bilgisi.runs[0]
        run_g_c_bilgisi.font.name = 'Comic Sans MS'
        run_g_c_bilgisi.font.size = Pt(14)

        # Değişiklikleri kaydedin
        diyet.save(str(ad_soyad) + ".docx")

        # EXCEL Dosyasına kaydetme:
        kisi_kontrol = str(sonraki_zaman.strftime("%m/%d/%Y"))
        # Verileri kayitlar.txt dosyasına ekle
        with open('kayitlar.txt', 'a') as file:
            file.write(f'{ad_soyad}\t{hafta}\t{kisi_kontrol}\n')
        print("Veriler kayitlar.txt dosyasına başarıyla eklendi.")


    except ValueError:
        result_text.set("Geçersiz değerler! Lütfen doğru değerleri girin.")

# Tkinter penceresini oluştur
window = tk.Tk()
window.title("BMI Hesaplayıcı")

# Etiketler ve Giriş Kutuları
label_ad_soyad = tk.Label(window, text="Adı ve Soyad:")
label_ad_soyad.grid(row=0, column=0, padx=10, pady=10)

entry_ad_soyad = tk.Entry(window)
entry_ad_soyad.grid(row=0, column=1, padx=10, pady=10)

label_ogun_sayisi = tk.Label(window, text="Kaç öğün:")
label_ogun_sayisi.grid(row=1, column=0, padx=10, pady=10)

entry_ogun_sayisi = tk.Entry(window)
entry_ogun_sayisi.grid(row=1, column=1, padx=10, pady=10)

label_kilo = tk.Label(window, text="Kilo (kg):")
label_kilo.grid(row=2, column=0, padx=10, pady=10)

entry_kilo = tk.Entry(window)
entry_kilo.grid(row=2, column=1, padx=10, pady=10)

label_boy = tk.Label(window, text="Boy (cm):")
label_boy.grid(row=3, column=0, padx=10, pady=10)

entry_boy = tk.Entry(window)
entry_boy.grid(row=3, column=1, padx=10, pady=10)

label_yas = tk.Label(window, text="Yaş:")
label_yas.grid(row=4, column=0, padx=10, pady=10)

entry_yas = tk.Entry(window)
entry_yas.grid(row=4, column=1, padx=10, pady=10)

# Hesapla düğmesi
calculate_button = tk.Button(window, text="Hesapla", command=calculate_bmi)
calculate_button.grid(row=5, column=0, columnspan=2, pady=10)

# BMI sonucunu gösteren etiket
result_text = tk.StringVar()
result_label = tk.Label(window, textvariable=result_text)
result_label.grid(row=6, column=0, columnspan=2, pady=10)

# Pencereyi başlat
window.mainloop()


